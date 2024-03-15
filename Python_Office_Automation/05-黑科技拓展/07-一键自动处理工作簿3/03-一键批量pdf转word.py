from pdfminer.pdfparser import PDFParser
from pdfminer.pdfparser import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
import os

def parse(fn, pdf_file_name):
    document = Document()

    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 提供初始化密码doc.initialize("lianxipython")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed

    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource, laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource, device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out, "get_text"):
                    # print(out.get_text(), type(out.get_text()))
                    content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                    # with open('test.txt','a') as f:
                    #     f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                    try:
                        document.add_paragraph(
                            content, style='ListBullet'  # 添加段落，样式为unordered list类型
                        )
                    except Exception as e:
                        print(e)
                document.save("../source_material/07/03一键批量pdf转word/word文件/" + pdf_file_name[:-4] + '.doc')  # 保存这个文档

# 获取指定后缀的文件，比如获取文件第07章：序列-字典中所有的.mp4文件。
def get_filename(path, filetype):  # 输入路径、文件类型 例如'.csv'
    file_name = []
    for root, dirs, files in os.walk(path):
        for i in files:
            if filetype + ' ' in i + ' ':  # 这里后面不加一个字母可能会出问题，加上一个（不一定是空格）可以解决99.99%的情况
                file_name.append(i)
    return file_name  # 输出由有后缀的文件名组成的列表

if __name__ == '__main__':

    path_file = "../source_material/07/03一键批量pdf转word/pdf文件/"
    file_name_list = get_filename(path_file, ".pdf")
    for pdf_file_name in file_name_list:
        fn = open(path_file + pdf_file_name, 'rb')
        parse(fn, pdf_file_name)
        print(pdf_file_name[:-4] + "已转word")
