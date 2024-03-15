# 导入库
from docx import Document

# 创建word文档
document = Document()

# 段落是word的基础。它们用于正文文本，但也用于标题和列标项目(如项目符号)
# 此方法返回对段落的引用
paragraph = document.add_paragraph("我是段落A")
paragraph = document.add_paragraph("我是段落B")

# 将段落插入在我上一个段落上面
prior_paragraph = paragraph.insert_paragraph_before("我会将段落插入在我上一个段落上面")

paragraph = document.add_paragraph("我是段落C")

# 这是一个分页符，用来断开页面
document.add_page_break()
paragraph = document.add_paragraph("因为上面有分页符，所以我们自动加入到下一页")


# 保存文档
document.save("../source_material/01/A/放假通知(1).docx")
print("执行完成")
