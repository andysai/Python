"""
对图片的操作
    # 导入库
    from docx import Document
    # 导入尺寸类
    from docx.shared import Inches
    # 导入位置类
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 创建word文档
    document = Document()

    # 添加图片:设定图片宽高比例
    # 设置宽，高会自适应，不管图片怎样，都会按照设定的宽度进行拉伸
    document.add_picture("../source_material/01/A/糟糕透了精彩极了.png", width=Inches(6))
"""

# 导入库
from docx import Document
# 导入尺寸类
from docx.shared import Inches
# 导入位置类
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建word文档
document = Document()

# 添加图片:设定图片宽高比例
# 设置宽，高会自适应，不管图片怎样，都会按照设定的宽度进行拉伸
# document.add_picture("../source_material/01/A/糟糕透了精彩极了.png", width=Inches(6))

# 添加图片:图片位置默认靠左，设定图片位置:LEFT CENTER RIGHT
image_p = document.add_paragraph()
image_run = image_p.add_run()
image_run.add_picture("../source_material/01/A/糟糕透了精彩极了.png", width=Inches(2))
image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
document.save("../source_material/01/A/放假通知(4).docx")
print("执行完成")
