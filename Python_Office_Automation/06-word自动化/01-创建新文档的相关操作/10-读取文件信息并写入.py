# 导入库
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
# 设置字体颜色
from docx.shared import RGBColor

# 创建word文档
document = Document()

def get_filename(path, filetype):
    file_name = []
    for filenames in os.walk(path):
        print(filenames[2])
        for filename in filenames[2]:
            if filetype in filename:
                file_name.append(filename)

    return file_name

# 写入文档
def read_write(file_name):
    # 创建word文档
    document = Document()

    # 设置全局字体
    document.styles["Normal"].font.name = u"宋体"
    document.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles["Normal"].font.size = Pt(15)
    document.add_heading("python简介", 0)

    for file in file_name:
        path = "D:/study/python/Python_Office_Automation/06-word自动化/source_material/01/text文档/" + file
        file_text = open(path, "r", encoding='utf-8')
        str_file = file_text.read()
        file_text.close()
        # 添加数据
        document.add_paragraph(str_file)

    # 保存文档
    document.save(r"D:\study\python\Python_Office_Automation\06-word自动化\source_material\01\text文档/合并文件.docx")

# python程序执行入口
if __name__ == "__main__":
    # 切换目录
    path = "../source_material/01"
    os.chdir(path)

    # 获取目录下的文本
    file_name = get_filename("./text文档", ".txt")

    # 读取文档文件并写入
    read_write(file_name)
