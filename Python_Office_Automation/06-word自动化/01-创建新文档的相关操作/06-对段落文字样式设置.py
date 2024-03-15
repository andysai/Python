# 导入库
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
# 设置字体颜色
from docx.shared import RGBColor

# 创建word文档
document = Document()

# 在一个段落中放置不同的字体样式
# 创建的段落对象
p1 = document.add_paragraph()
text1 = p1.add_run("我是中文字体wo shi zhong wen zi ti")
text2 = p1.add_run("我是中文字体wo shi zhong wen zi ti")

# 针对某个字体对象进行设定字体
# 设定字体大小
text1.font.size = Pt(15)
# 设定字体是否加粗
text1.bold = True
# 设定字体是否加下划线
text1.underline = True
# 设定字体颜色
text1.font.color.rgb = RGBColor(250, 25, 25)

# 保存文档
document.save("../source_material/01/A/放假通知(5).docx")
print("执行完成")
