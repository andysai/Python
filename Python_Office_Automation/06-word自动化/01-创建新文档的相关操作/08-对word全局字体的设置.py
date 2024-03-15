# 导入库
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
# 设置字体颜色
from docx.shared import RGBColor

# 创建word文档
document = Document()

# 设置全局字体
document.styles["Normal"].font.name = "宋体"
document.styles["Normal"].font.size = Pt(20)
document.styles["Normal"].font.color.rgb = RGBColor(255, 25, 25)
document.styles["Normal"].font.underline = True
document.styles["Normal"].font.bold = True
document.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 添加段落
p1 = document.add_paragraph("君不见黄河之水天上来，奔流到海不复回。")
p2 = document.add_paragraph("举头望明月，低头思故乡。")

# 保存文档
document.save("../source_material/01/A/放假通知(7).docx")
print("执行完成")
