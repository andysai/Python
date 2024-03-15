# 导入库
from docx import Document

# 创建word文档
document = Document()
# 加载原有后缀为docx的文档

# path可以给绝对路径或者相对路径
# word = docx.Document("../s")

# 1 如果出现标题添加标题问题，卸载重装python-docx
# 2 添加标题:标题段落包含文本，段落样式级别决定，从0到9级
# 省略不写默认1级
document.add_heading("我是默认1级标题")
document.add_heading("我是0级标题", 0)
document.add_heading("我是1级标题", 1)
document.add_heading("我是2级标题", 2)
document.add_heading("我是3级标题", 3)

# 保存文档
document.save("../source_material/01/A/放假通知.docx")
print("执行完成")
