# 导入库
from docx import Document

# 创建word文档
document = Document()

article = {
    "title": "精彩极了，糟糕透了",
    "content": ['记得七八岁的时候，我写了一首诗。母亲一念完那首诗，眼睛亮亮地，兴奋地嚷道：“巴迪，这真是你写的字吗）多美的诗啊！精彩极了！”她搂住了我，赞扬声雨点般地落在我上。我既腼腆又得意洋洋，点头告诉她这首诗确实是我写的。她高兴得再次拥抱我。',
                '“妈妈，爸爸下午什么时候回来？”我红着脸问。我有点迫不及待，想立刻让父亲看看我写的诗。“他晚上七点钟回来。”母亲摸着我的脑袋，笑着说。',
                '整个下午，我用最漂亮的花体字把诗认认真真地重新誊写了一遍，还用彩色笔在它周围描上了一圈花边。将近七点钟的时候，我悄悄走进饭厅，满怀信心地把它平平整整地放在餐桌父亲的位置上。']
}

document.add_heading(article["title"], 0)
for i in article["content"]:
    paragraph = document.add_paragraph(i)

# 保存文档
document.save("../source_material/01/A/放假通知(2).docx")
print("执行完成")
