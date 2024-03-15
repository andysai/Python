# 导入模块
from docx import Document
import os
from docx.oxml.ns import qn
from docx.shared import Pt

# 获取指定后缀的文件
def get_filename(path, filetype):
    file_name = []
    for root, dirs, files in os.walk(path):
        for i in files:
            if filetype+'' in i+'':
                file_name.append(i)
    return file_name

# 写入数据
def input_value(path, file_name_list):
    # 创建word文档
    document_new = Document()

    # 设置全局字体
    document_new.styles['Normal'].font.name = "宋体"
    document_new.styles['Normal'].font.size = Pt(20)
    document_new.styles['Normal'].font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    for file_name in file_name_list:
        document = Document(path+"/"+file_name)
        for paragraph in document.paragraphs:
            document_new.add_paragraph(paragraph.text)

    #保存文档
    document_new.save("../source_material/04/02一键合并word文档/合并文档.docx")
    print("执行完毕")

# 程序执行入口
if __name__ == '__main__':
    path = "../source_material/04/02一键合并word文档/要合并的word文档"
    # 获取文件夹下所有的文件
    file_name_list = get_filename(path, ".docx")
    # 写入数据
    input_value(path, file_name_list)
