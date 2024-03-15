# 导入模块
from docx import Document
import re
from docx.oxml.ns import qn
from docx.shared import Pt

# 加载原文档
document = Document("../source_material/04/01一键拆分word文档/李白古诗大全.docx")

# 读取段落信息
paragraph_list = []
for paragraph in document.paragraphs:
    if len(paragraph.text) != 0:
        paragraph_list.append(paragraph.text.strip())

# 将下标放进去
index_value_list = []
for paragraph_value in paragraph_list:
    if paragraph_value[0].isdigit():
        index_value_list.append(paragraph_list.index(paragraph_value))

for index_num in index_value_list:
    if index_value_list.index(index_num)+1 < len(index_value_list):
        list_index = index_value_list[index_value_list.index(index_num)+1]

        # 创建word文档
        document = Document()

        # 设置全局字体
        document.styles['Normal'].font.name = "宋体"
        document.styles['Normal'].font.size = Pt(20)
        document.styles['Normal'].font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        file_names = ""

        for i in range(index_num, list_index):
            # 判断每一段第一个字符是否为数字
            if paragraph_list[i][0].isdigit():
                file_names = re.sub(r"[0-9]{1,2}|、", "", paragraph_list[i])
                document.add_heading(file_names, 0)
            else:
                paragraph = document.add_paragraph(paragraph_list[i])

        document.save("../source_material/04/01一键拆分word文档/李白古诗/" + file_names + ".docx")
        print(file_names + "正在保存中...")
